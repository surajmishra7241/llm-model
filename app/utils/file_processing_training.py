import os
import logging
import requests
from typing import Optional, List
from io import BytesIO
import tempfile
from pathlib import Path

# Document processing
import PyPDF2
import docx
from bs4 import BeautifulSoup
import json
import csv

# Image processing (OCR)
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR dependencies not available. Install Pillow and pytesseract for image processing.")

# Audio processing (for future use)
try:
    import speech_recognition as sr
    SPEECH_AVAILABLE = True
except ImportError:
    SPEECH_AVAILABLE = False
    logging.warning("Speech recognition not available. Install SpeechRecognition for audio processing.")

logger = logging.getLogger(__name__)

def process_document(filename: str, content: bytes) -> Optional[str]:
    """Process various document types and extract text"""
    try:
        file_ext = os.path.splitext(filename.lower())[1]
        
        if file_ext == '.pdf':
            return extract_text_from_pdf(content)
        elif file_ext in ['.doc', '.docx']:
            return extract_text_from_word(content)
        elif file_ext in ['.txt', '.md']:
            return content.decode('utf-8', errors='ignore')
        elif file_ext in ['.html', '.htm']:
            return extract_text_from_html(content)
        elif file_ext == '.json':
            return extract_text_from_json(content)
        elif file_ext == '.csv':
            return extract_text_from_csv(content)
        elif file_ext == '.xml':
            return extract_text_from_xml(content)
        elif file_ext == '.rtf':
            return extract_text_from_rtf(content)
        else:
            # Try to decode as text for unknown formats
            try:
                return content.decode('utf-8', errors='ignore')
            except:
                logger.warning(f"Unknown file format: {file_ext}")
                return None
                
    except Exception as e:
        logger.error(f"Error processing document {filename}: {str(e)}")
        return None

def process_image(filename: str, content: bytes) -> Optional[str]:
    """Process images using OCR to extract text"""
    if not OCR_AVAILABLE:
        logger.warning("OCR not available, skipping image processing")
        return None
    
    try:
        # Open image from bytes
        image = Image.open(BytesIO(content))
        
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Extract text using OCR
        text = pytesseract.image_to_string(image)
        
        # Clean up extracted text
        text = text.strip()
        if len(text) < 10:  # Minimum text threshold
            logger.warning(f"Very little text extracted from image: {filename}")
            return None
        
        return text
        
    except Exception as e:
        logger.error(f"Error processing image {filename}: {str(e)}")
        return None

def extract_text_from_pdf(content: bytes) -> Optional[str]:
    """Extract text from PDF content"""
    try:
        pdf_file = BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return None

def extract_text_from_word(content: bytes) -> Optional[str]:
    """Extract text from Word documents"""
    try:
        doc_file = BytesIO(content)
        doc = docx.Document(doc_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from Word document: {str(e)}")
        return None

def extract_text_from_html(content: bytes) -> Optional[str]:
    """Extract text from HTML content"""
    try:
        html_content = content.decode('utf-8', errors='ignore')
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text and clean it up
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text from HTML: {str(e)}")
        return None

def extract_text_from_json(content: bytes) -> Optional[str]:
    """Extract text from JSON content"""
    try:
        json_content = content.decode('utf-8', errors='ignore')
        data = json.loads(json_content)
        
        def extract_strings(obj, strings_list):
            """Recursively extract all string values from JSON"""
            if isinstance(obj, dict):
                for value in obj.values():
                    extract_strings(value, strings_list)
            elif isinstance(obj, list):
                for item in obj:
                    extract_strings(item, strings_list)
            elif isinstance(obj, str) and len(obj.strip()) > 2:
                strings_list.append(obj.strip())
        
        strings = []
        extract_strings(data, strings)
        
        return "\n".join(strings)
        
    except Exception as e:
        logger.error(f"Error extracting text from JSON: {str(e)}")
        return None

def extract_text_from_csv(content: bytes) -> Optional[str]:
    """Extract text from CSV content"""
    try:
        csv_content = content.decode('utf-8', errors='ignore')
        csv_file = BytesIO(csv_content.encode())
        
        # Try different delimiters
        for delimiter in [',', ';', '\t', '|']:
            try:
                csv_file.seek(0)
                reader = csv.reader(csv_content.splitlines(), delimiter=delimiter)
                rows = list(reader)
                
                if len(rows) > 0 and len(rows[0]) > 1:
                    # Convert CSV to readable text
                    text_lines = []
                    headers = rows[0] if rows else []
                    
                    for row in rows:
                        if headers and len(row) == len(headers):
                            row_text = []
                            for header, value in zip(headers, row):
                                if value.strip():
                                    row_text.append(f"{header}: {value}")
                            if row_text:
                                text_lines.append(", ".join(row_text))
                        else:
                            text_lines.append(", ".join(cell for cell in row if cell.strip()))
                    
                    return "\n".join(text_lines)
            except:
                continue
        
        # Fallback: return raw content
        return csv_content
        
    except Exception as e:
        logger.error(f"Error extracting text from CSV: {str(e)}")
        return None

def extract_text_from_xml(content: bytes) -> Optional[str]:
    """Extract text from XML content"""
    try:
        xml_content = content.decode('utf-8', errors='ignore')
        soup = BeautifulSoup(xml_content, 'xml')
        
        # Extract all text content
        text = soup.get_text(separator=' ', strip=True)
        
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text from XML: {str(e)}")
        return None

def extract_text_from_rtf(content: bytes) -> Optional[str]:
    """Extract text from RTF content (basic implementation)"""
    try:
        rtf_content = content.decode('utf-8', errors='ignore')
        
        # Very basic RTF parsing - remove RTF control codes
        import re
        
        # Remove RTF header
        text = re.sub(r'\\rtf\d+.*?(?=\\)', '', rtf_content)
        
        # Remove RTF control words
        text = re.sub(r'\\[a-z]+\d*\s?', '', text)
        
        # Remove curly braces
        text = re.sub(r'[{}]', '', text)
        
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text if text else None
        
    except Exception as e:
        logger.error(f"Error extracting text from RTF: {str(e)}")
        return None

def extract_text_from_url(url: str) -> Optional[str]:
    """Extract text content from URL"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        content_type = response.headers.get('content-type', '').lower()
        
        if 'text/html' in content_type:
            return extract_text_from_html(response.content)
        elif 'application/json' in content_type:
            return extract_text_from_json(response.content)
        elif 'text/plain' in content_type:
            return response.text
        elif 'application/pdf' in content_type:
            return extract_text_from_pdf(response.content)
        else:
            # Try to extract as text
            try:
                return response.text
            except:
                logger.warning(f"Unknown content type for URL {url}: {content_type}")
                return None
                
    except Exception as e:
        logger.error(f"Error extracting text from URL {url}: {str(e)}")
        return None

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks"""
    if not text or len(text.strip()) < 10:
        return []
    
    text = text.strip()
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        if end >= len(text):
            chunks.append(text[start:])
            break
        
        # Try to break at sentence boundary
        chunk = text[start:end]
        last_period = chunk.rfind('.')
        last_newline = chunk.rfind('\n')
        last_space = chunk.rfind(' ')
        
        # Choose the best breaking point
        break_point = max(last_period, last_newline, last_space)
        if break_point > start + chunk_size // 2:  # Only if break point is not too early
            actual_end = start + break_point + 1
        else:
            actual_end = end
        
        chunks.append(text[start:actual_end].strip())
        start = actual_end - overlap
        
        # Ensure we make progress
        if start <= 0:
            start = actual_end
    
    return [chunk for chunk in chunks if len(chunk.strip()) > 10]

def process_audio_file(filename: str, content: bytes) -> Optional[str]:
    """Process audio files using speech recognition (future feature)"""
    if not SPEECH_AVAILABLE:
        logger.warning("Speech recognition not available")
        return None
    
    try:
        # Save audio to temporary file
        with tempfile.NamedTemporaryFile(suffix=os.path.splitext(filename)[1], delete=False) as temp_file:
            temp_file.write(content)
            temp_path = temp_file.name
        
        try:
            # Initialize recognizer
            r = sr.Recognizer()
            
            # Load audio file
            with sr.AudioFile(temp_path) as source:
                audio = r.record(source)
            
            # Recognize speech
            text = r.recognize_google(audio)
            return text
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.unlink(temp_path)
                
    except Exception as e:
        logger.error(f"Error processing audio file {filename}: {str(e)}")
        return None

# File validation utilities
def validate_file_size(content: bytes, max_size_mb: int = 50) -> bool:
    """Validate file size"""
    size_mb = len(content) / (1024 * 1024)
    return size_mb <= max_size_mb

def validate_file_type(filename: str, allowed_extensions: set = None) -> bool:
    """Validate file type by extension"""
    if allowed_extensions is None:
        from app.models.training_model import SUPPORTED_EXTENSIONS
        allowed_extensions = SUPPORTED_EXTENSIONS
    
    ext = os.path.splitext(filename.lower())[1]
    return ext in allowed_extensions

def get_file_info(filename: str, content: bytes) -> dict:
    """Get file information"""
    return {
        "filename": filename,
        "extension": os.path.splitext(filename.lower())[1],
        "size_bytes": len(content),
        "size_mb": round(len(content) / (1024 * 1024), 2),
        "estimated_text_length": len(content) // 2,  # Rough estimate
    }